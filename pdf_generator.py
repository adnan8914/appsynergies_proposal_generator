import streamlit as st
import datetime
import tempfile
import os
from docx import Document
import platform

# Check if running on Windows or cloud
IS_WINDOWS = platform.system() == "Windows"

if IS_WINDOWS:
    import win32com.client
    import pythoncom
    from docx2pdf import convert
else:
    # Alternative PDF conversion for cloud (could offer DOCX download only)
    def convert(input_docx, output_pdf):
        st.warning("PDF conversion is only available in Windows. Downloading DOCX instead.")
        return False

def convert(input_docx, output_pdf):
    if IS_WINDOWS:
        try:
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch('Word.Application')
            doc = None
            try:
                word.Visible = False
                
                # Ensure paths are absolute and use raw strings
                input_path = os.path.abspath(input_docx)
                output_path = os.path.abspath(output_pdf)
                
                # Open document
                doc = word.Documents.Open(input_path)
                
                # Try different SaveAs methods
                try:
                    doc.SaveAs(FileName=output_path, FileFormat=17)
                except:
                    try:
                        doc.SaveAs(output_path, 17)
                    except:
                        doc.ExportAsFixedFormat(output_path, 17)
                return True
            finally:
                # Clean up
                if doc:
                    doc.Close()
                word.Quit()
        except Exception as e:
            st.error(f"Conversion error: {str(e)}")
            return False
        finally:
            pythoncom.CoUninitialize()
    else:
        return convert_to_pdf_cloud(input_docx, output_pdf)

def convert_to_pdf_cloud(input_docx, output_pdf):
    try:
        import subprocess
        # Convert using LibreOffice
        cmd = ['libreoffice', '--headless', '--convert-to', 'pdf', input_docx, '--outdir', os.path.dirname(output_pdf)]
        subprocess.run(cmd, timeout=30)
        return True
    except Exception as e:
        st.error(f"PDF Conversion error: {str(e)}")
        return False

# Define the template path
TEMPLATE_PATH = r"C:\Users\Adnan\.cursor-tutor\new_as_pdf_generator\Ai_automation.docx"

def replace_text_preserve_formatting(doc, replacements):
    """Replace text while preserving formatting and images"""
    def replace_in_paragraph(paragraph, replacements):
        paragraph_text = paragraph.text
        runs = paragraph.runs
        
        # Debug print to see what we're dealing with
        if "Additional Features" in paragraph_text:
            st.write("Found Additional Features paragraph:")
            st.write("Runs:", [r.text for r in runs])
        
        for key, value in replacements.items():
            if key in paragraph_text:
                # Format price values if needed
                if "price" in key.lower() or "amount" in key.lower() or key == "{Additional}":
                    if isinstance(value, (int, float)):
                        value = f"$ {value:,.2f}"
                
                # Special handling for Additional Features text box
                if "Additional Features" in paragraph_text and "Enhancements" in paragraph_text:
                    # Try to find the exact run with the placeholder
                    for i, run in enumerate(runs):
                        if "{Additional}" in run.text:
                            # Count tabs/spaces before the placeholder
                            prefix = run.text[:run.text.find("{Additional}")]
                            run.text = prefix + str(value)
                            # Clear any remaining parts
                            for j in range(i + 1, len(runs)):
                                if "{" in runs[j].text or "}" in runs[j].text:
                                    runs[j].text = ""
                            return
                
                # Normal replacement logic for other fields
                key_runs = []
                start_index = -1
                current_key_part = ""
                
                for i, run in enumerate(runs):
                    if not run.text:
                        continue
                    
                    current_key_part += run.text
                    if key in current_key_part:
                        for j in range(start_index if start_index >= 0 else i, i + 1):
                            key_runs.append(runs[j])
                        start_index = -1
                        current_key_part = ""
                    elif any(part in current_key_part for part in ["{", key[1:len(key)-1]]):
                        if start_index < 0:
                            start_index = i
                    else:
                        start_index = -1
                        current_key_part = ""
                
                if key_runs:
                    key_runs[0].text = str(value)
                    for run in key_runs[1:]:
                        run.text = ""

    # Process all paragraphs including those in text boxes
    for paragraph in doc.element.xpath('//w:p'):
        try:
            p = doc.paragraphs[0].__class__(paragraph, doc.paragraphs[0]._parent)
            if p.text.strip():
                replace_in_paragraph(p, replacements)
        except:
            pass

    # Process regular paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            replace_in_paragraph(paragraph, replacements)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        replace_in_paragraph(paragraph, replacements)

    # Process text boxes using Word XML
    for element in doc._element.xpath('.//w:drawing//wp:anchor//w:txbxContent//w:p'):
        try:
            text = element.text
            for key, value in replacements.items():
                if key in text:
                    if "price" in key.lower() or "amount" in key.lower() or key == "{Additional}":
                        if isinstance(value, (int, float)):
                            value = f"$ {value:,.2f}"
                    element.text = text.replace(key, str(value))
        except:
            continue

    # Try to process shapes directly
    try:
        for shape in doc.inline_shapes:
            if shape._inline.graphic.graphicData.pic is not None:
                txbox = shape._inline.graphic.graphicData.pic.nvPicPr.cNvPr.txBox
                if txbox is not None:
                    for paragraph in txbox.paragraphs:
                        replace_in_paragraph(paragraph, replacements)
    except:
        pass

def main():
    st.title("Proposal Generator")
    
    # Proposal Type Selection
    proposal_type = st.selectbox(
        "Select Proposal Type",
        ["AI Automation Proposal", "Digital Marketing Proposal", "Business Automations Proposal", "IT Consultation Contract"]
    )
    
    if proposal_type == "AI Automation Proposal":
        template_path = r"C:\Users\Adnan\.cursor-tutor\new_as_pdf_generator\Ai_automation.docx"
        
        # Client Information
        st.header("Client Information")
        col1, col2 = st.columns(2)
        with col1:
            client_name = st.text_input("Client Name", key="ai_name")
            email = st.text_input("Email", key="ai_email")
            phone = st.text_input("Phone", key="ai_phone")
        with col2:    
            country = st.text_input("Country", key="ai_country")
            proposal_date = st.date_input("Date", datetime.datetime.now(), key="ai_date")
        
        # Project Pricing
        st.header("Project Pricing")
        col1, col2 = st.columns(2)
        with col1:
            landing_page_price = st.number_input("Landing Page Website", min_value=0.0, step=0.01, format="%.2f")
            admin_panel_price = st.number_input("Admin Panel", min_value=0.0, step=0.01, format="%.2f")
            crm_price = st.number_input("CRM Automations", min_value=0.0, step=0.01, format="%.2f")
        with col2:
            manychat_price = st.number_input("ManyChat & Make Automation", min_value=0.0, step=0.01, format="%.2f")
            social_media_price = st.number_input("Social Media Automation", min_value=0.0, step=0.01, format="%.2f")
            ai_calling_price = st.number_input("AI Calling", min_value=0.0, step=0.01, format="%.2f")
        
        # Additional Features & Enhancements
        additional_features_price = st.number_input("Additional Features & Enhancements (USD per week)", 
                                                  min_value=0.0, 
                                                  step=0.01, 
                                                  format="%.2f",
                                                  value=250.00)  # Default value set to 250
        
        total_price = (landing_page_price + admin_panel_price + crm_price + 
                      manychat_price + social_media_price + ai_calling_price)
        annual_maintenance = total_price * 0.20
        
        # Signature Details
        st.header("Signature Details")
        company_representative = st.text_input("Company Representative")

        # Updated replacements dictionary
        replacements = {
            "{client_name}": client_name,
            "{Email_address}": email,
            "{Phone_no}": phone,
            "{country_name}": country,
            "{date}": proposal_date.strftime("%d/%m/%Y"),
            "{landing page price}": landing_page_price,
            "{admin panel price}": admin_panel_price,
            "{CRM Automation price}": crm_price,
            "{Manychat price}": manychat_price,
            "{SMP price}": social_media_price,
            "{AI calling price}": ai_calling_price,
            "{Total amount}": total_price,
            "{AM price}": annual_maintenance,
            "{Additional}": additional_features_price,
            "{company_representative}": company_representative,
        }

    elif proposal_type == "Digital Marketing Proposal":
        template_path = r"C:\Users\Adnan\.cursor-tutor\new_as_pdf_generator\DM Proposal.docx"
        
        # Digital Marketing proposal inputs
        st.header("Client Information")
        col1, col2 = st.columns(2)
        with col1:
            client_name = st.text_input("Client Name", key="dm_name")
            designation = st.text_input("Designation", key="dm_designation")
            contact_no = st.text_input("Contact Number", key="dm_contact")
        with col2:
            email_id = st.text_input("Email ID", key="dm_email")
            proposal_date = st.date_input("Date", datetime.datetime.now(), key="dm_date")
        
        # Mutually Agreed Points
        mutually_agreed_points = st.text_area("Mutually Agreed Points")
        
        # Digital Marketing Services Pricing
        st.header("Digital Marketing Services Pricing")
        col1, col2 = st.columns(2)
        with col1:
            social_media_posts = st.number_input("30 Creative Social Media Posts", min_value=0.0, step=0.01, format="%.2f", key="dm_smp")
            research_and_dev = st.number_input("Marketing Research + 1 Month Ads", min_value=0.0, step=0.01, format="%.2f", key="dm_rnd")
        with col2:
            monthly_cost = st.number_input("Monthly Maintenance", min_value=0.0, step=0.01, format="%.2f", key="dm_monthly")
        
        # Calculate totals
        subtotal = social_media_posts + research_and_dev + monthly_cost
        gst = subtotal * 0.18  # 18% GST
        total_amount = subtotal + gst
        
        # Calculate payment schedule with dollar formatting
        advance = total_amount * 0.5
        balance = total_amount * 0.5
        
        replacements = {
            "{client_name}": client_name,
            "{designation}": designation,
            "{contact_no}": contact_no,
            "{email_id}": email_id,
            "{date}": proposal_date.strftime("%d/%m/%Y"),
            "{Mutually_agreed_points}": mutually_agreed_points,
            "{3d_SMP}": f"$ {social_media_posts:,.2f}",
            "{R&D}": f"$ {research_and_dev:,.2f}",
            "{monthly_cost}": f"$ {monthly_cost:,.2f}",
            "{gst}": f"$ {gst:,.2f}",
            "{total_amount}": f"$ {total_amount:,.2f}",
            "{Advance}": f"$ {advance:,.2f}",
            "{balance}": f"$ {balance:,.2f}"
        }

    elif proposal_type == "Business Automations Proposal":
        template_path = r"C:\Users\Adnan\.cursor-tutor\new_as_pdf_generator\Business Automations Proposal.docx"
        
        # Client Information
        st.header("Client Information")
        col1, col2 = st.columns(2)
        with col1:
            client_name = st.text_input("Client Name", key="ba_name")
            contact_no = st.text_input("Contact Number", key="ba_contact")
            email_id = st.text_input("Email ID", key="ba_email_id")
        with col2:
            proposal_date = st.date_input("Date", datetime.datetime.now(), key="ba_date")
            validity_date = st.date_input("Validity Date", key="ba_validity")

        # Mutually Agreed Points
        mutually_agreed_points = st.text_area("Mutually Agreed Points", key="ba_points")
        
        # Week 1 Details
        st.header("Week 1 Details")
        week1_descrptn = st.text_area("Week 1 Description", key="ba_week1_desc")
        week1_price = st.number_input("Week 1 Price", min_value=0.0, step=0.01, format="%.2f", key="ba_week1_price")

        # Future Services Pricing
        st.header("Future Services Pricing")
        col1, col2 = st.columns(2)
        with col1:
            ai_auto_price = st.number_input("AI Automations Price", min_value=0.0, step=0.01, format="%.2f", key="ba_ai_auto")
            whts_price = st.number_input("WhatsApp Automation Price", min_value=0.0, step=0.01, format="%.2f", key="ba_whts")
            crm_price = st.number_input("CRM Setup Price", min_value=0.0, step=0.01, format="%.2f", key="ba_crm")
            email_price = st.number_input("Email Marketing Setup Price", min_value=0.0, step=0.01, format="%.2f", key="ba_email_price")
            make_price = st.number_input("Make/Zapier Automation Price", min_value=0.0, step=0.01, format="%.2f", key="ba_make")
        with col2:
            firefly_price = st.number_input("Firefly Meeting Price", min_value=0.0, step=0.01, format="%.2f", key="ba_firefly")
            chatbot_price = st.number_input("AI Chatbot Price", min_value=0.0, step=0.01, format="%.2f", key="ba_chatbot")
            pdf_gen_pr = st.number_input("PDF Generation Price", min_value=0.0, step=0.01, format="%.2f", key="ba_pdf")
            ai_mdl_price = st.number_input("AI Social Media Price", min_value=0.0, step=0.01, format="%.2f", key="ba_ai_mdl")
            cstm_ai_price = st.number_input("Custom AI Models Price", min_value=0.0, step=0.01, format="%.2f", key="ba_cstm_ai")

        # Replacements dictionary
        replacements = {
            "{client_name}": client_name,
            "{contact_no}": contact_no,
            "{email_id}": email_id,
            "{date}": proposal_date.strftime("%d/%m/%Y"),
            "{validity_date}": validity_date.strftime("%d/%m/%Y"),
            "{mutually_agreed_points}": mutually_agreed_points,
            "{week1_descrptn}": week1_descrptn,
            "{week1_price}": f"$ {week1_price:,.2f}",
            "{ai_auto_price}": f"$ {ai_auto_price:,.2f}",
            "{whts_price}": f"$ {whts_price:,.2f}",
            "{crm_price}": f"$ {crm_price:,.2f}",
            "{email_price}": f"$ {email_price:,.2f}",
            "{make_price}": f"$ {make_price:,.2f}",
            "{firefly_price}": f"$ {firefly_price:,.2f}",
            "{chatbot_price}": f"$ {chatbot_price:,.2f}",
            "{pdf_gen_pr}": f"$ {pdf_gen_pr:,.2f}",
            "{ai_mdl_price}": f"$ {ai_mdl_price:,.2f}",
            "{cstm_ai_price}": f"$ {cstm_ai_price:,.2f}"
        }

    else:  # IT Consultation Contract
        template_path = r"C:\Users\Adnan\.cursor-tutor\new_as_pdf_generator\Contract Agreement.docx"
        
        # Client Information
        st.header("Contract Information")
        col1, col2 = st.columns(2)
        with col1:
            client_name = st.text_input("Client Name", key="contract_name")
            client_company_address = st.text_area("Company Address", key="contract_address")
        with col2:
            contract_date = st.date_input("Contract Date", datetime.datetime.now(), key="contract_date")

        # Replacements dictionary
        replacements = {
            "{date}": contract_date.strftime("%d/%m/%Y"),
            "{client_name}": client_name,
            "{client_company_address}": client_company_address
        }

    # Generate proposal button and processing - same for all types
    if st.button("Generate Proposal"):
        if proposal_type == "AI Automation Proposal":
            if not all([client_name, email, phone, country]):
                st.error("Please fill in all required fields")
                return
        elif proposal_type == "Digital Marketing Proposal":
            if not all([client_name, designation, contact_no, email_id]):
                st.error("Please fill in all required fields")
                return
        elif proposal_type == "Business Automations Proposal":
            if not all([client_name, contact_no, email_id]):
                st.error("Please fill in all required fields")
                return
        else:  # IT Consultation Contract
            if not all([client_name, client_company_address]):
                st.error("Please fill in all required fields")
                return
            
        try:
            doc = Document(template_path)
            replace_text_preserve_formatting(doc, replacements)
            
            temp_dir = tempfile.mkdtemp()
            output_docx = os.path.join(temp_dir, f"{proposal_type}_{client_name}.docx")
            doc.save(output_docx)
            
            if IS_WINDOWS:
                # Windows conversion
                output_pdf = os.path.join(temp_dir, f"{proposal_type}_{client_name}.pdf")
                if convert(output_docx, output_pdf):
                    with open(output_pdf, "rb") as pdf_file:
                        pdf_bytes = pdf_file.read()
                        st.download_button(
                            label="Download Proposal PDF",
                            data=pdf_bytes,
                            file_name=f"{proposal_type}_{client_name}.pdf",
                            mime="application/pdf"
                        )
                    st.success("Proposal generated successfully!")
            else:
                # Cloud conversion
                output_pdf = os.path.join(temp_dir, f"{proposal_type}_{client_name}.pdf")
                if convert_to_pdf_cloud(output_docx, output_pdf):
                    with open(output_pdf, "rb") as pdf_file:
                        pdf_bytes = pdf_file.read()
                        st.download_button(
                            label="Download Proposal PDF",
                            data=pdf_bytes,
                            file_name=f"{proposal_type}_{client_name}.pdf",
                            mime="application/pdf"
                        )
                    st.success("Proposal generated successfully!")
            
        except Exception as e:
            st.error(f"Error generating proposal: {str(e)}")
        finally:
            try:
                os.rmdir(temp_dir)
            except:
                pass

if __name__ == "__main__":
    main()